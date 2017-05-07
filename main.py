from docx import Document
from openpyxl import load_workbook
from openpyxl.styles import PatternFill
import re

LP_REGEX = re.compile('\d+')
SUM_REGEX = re.compile('(\d+(\.|\s))*\d+,\d+')
COLOR_YELLOW = 'FFFF00'
COLOR_ORANGE = 'FF7700'


def delete_paragraph(paragraph):
    """Function copied from creator of python-docx.
    For now there is no other way to delete a paragraph"""
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None


def replace_text_in_cell(cell, text):
    """Select first run from first paragraph, 
    replace text in it and delete all other runs and paragraphs"""
    for paragraph in cell.paragraphs[1:]:
        delete_paragraph(paragraph)
    cell.paragraphs[0].runs[0].text = text
    del cell.paragraphs[0].runs[1:]


def get_lp_format(document):
    """Return first found format with digits in form tables.
    If not found, return default format"""

    for table in document.tables:
        lp = table.cell(1, 0).text
        if LP_REGEX.search(lp):
            return lp
    else:
        return '\n1.'


def clean_value(value):
    """Clean value extracted from docx or xlsx, 
    delete all formatting and convert to float."""
    if type(value) not in (float, int):
        try:
            value = SUM_REGEX.search(value).group(0)
        except AttributeError:
            return 0
        value = value.replace('.', '')
        value = value.replace(' ', '')
        value = float(value.replace(',', '.'))
    return value


def colour_range_cells(ws, start, end, color):
    """Set background of all cells in range to given color"""
    cell_range = start + ':' + end
    rows = ws[cell_range]
    for row in rows:
        for cell in row:
            cell.fill = PatternFill("solid", fgColor=color)


# Open word document and extract tables
document = Document('documents/word.docx')
tables = document.tables

# Open spreadsheet document and extract actual rows(as cells are differently merged)
workbook = load_workbook('documents/excel.xlsx', data_only=True)
spreadsheet = workbook.active
merged_rows = [(y[0][1:], y[1][1:]) for y in
               [x.split(":") for x in spreadsheet.merged_cell_ranges if 'A' in x]]

# Get default format of lp from word document in case of for example empty cell
lp_format = get_lp_format(document)

for index, table in enumerate(tables):
    # Insert correct lp.
    lp = table.cell(1, 0).text
    if LP_REGEX.search(lp):
        lp = LP_REGEX.sub(str(index + 1), lp)
    else:
        lp = LP_REGEX.sub(str(index + 1), lp_format)
    replace_text_in_cell(table.cell(1, 0), lp)

    # Extract sum value from docx document
    # Weird -1 index is necessary because of how word interprets tables
    document_value = table.row_cells(2)[-1].text

    # Clean document value
    document_value = clean_value(document_value)

    # Extract sum value from xlsx
    spreadsheet_value = spreadsheet['J' + merged_rows[index][1]].value

    # Clean spreadsheet value
    spreadsheet_value = clean_value(spreadsheet_value)

    # Report if something's wrong with spreadsheet value
    if spreadsheet_value == 0:
        print("Check sum value in spreadsheet.\n")

    # Check if document value and spreadsheet value are the same and
    # if not copy spreadsheet value to document value
    if document_value != round(spreadsheet_value, 2):
        print(str(index + 1), "Document and spreadsheet values are not the same: (" + str(document_value) + ' ' +str(spreadsheet_value)
              + "). Copying value from spreadsheet.\n")
        document_value = spreadsheet_value
        value_to_replace = ','.join('{0:,.2f}'.format(document_value).replace(',', '.').rsplit('.', 1))
        text_to_replace = SUM_REGEX.sub(value_to_replace, table.row_cells(2)[-1].text)
        replace_text_in_cell(table.row_cells(2)[-1], text_to_replace)

    # Get second sum from document, clean it and check if equals first. If not, replace it.
    document_value2 = table.row_cells(3)[-1].text
    document_value2 = clean_value(document_value2)

    if document_value != document_value2:
        print(str(index + 1), "Sums for row 2 and 3 in document are not the same: (" + str(document_value) + ' ' + str(document_value2)
              + "). Copying sum from row 2.\n")
        value_to_replace = ','.join('{0:,.2f}'.format(document_value).replace(',', '.').rsplit('.', 1))
        text_to_replace = SUM_REGEX.sub(value_to_replace, table.row_cells(3)[-1].text)
        replace_text_in_cell(table.row_cells(3)[-1], text_to_replace)

    # Extract creditor group, check if sum is higher than 7,000 and check if creditor group is correct.
    # Then colour rows in spreadsheet in yellow for group I and orange for group II.
    creditor_group = table.row_cells(8)[-1].text
    if document_value > 7000:
        if creditor_group != 'II':
            creditor_group = 'II'
            replace_text_in_cell(table.row_cells(8)[-1], creditor_group)
            print(str(index + 1), "Group II set.\n")
        colour_range_cells(spreadsheet, ('A' + merged_rows[index][0]),('R' + merged_rows[index][1]), COLOR_YELLOW)
    else:
        if creditor_group != 'I':
            creditor_group = 'I'
            replace_text_in_cell(table.row_cells(8)[-1], creditor_group)
            print(str(index + 1), "Group I set.\n")
        colour_range_cells(spreadsheet, ('A' + merged_rows[index][0]), ('R' + merged_rows[index][1]), COLOR_ORANGE)

document.save('documents/new.docx')
workbook.save('documents/new.xlsx')